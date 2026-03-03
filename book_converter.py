
import os
import re
from abc import ABC, abstractmethod

from pdf_processing import (
    extract_pdf_chapters,
    finalize_pdf_processing,
    prepare_pdf_for_processing,
)

class Book(ABC):
    """Abstract base class for a book."""

    def __init__(self, filepath, **options):
        self.filepath = filepath
        self.options = options

    @abstractmethod
    def extract_chapters(self):
        """Extract chapters from the book."""
        pass

class EpubBook(Book):
    """A book in EPUB format."""

    def extract_chapters(self):
        try:
            from ebooklib import epub
        except ImportError as exc:
            raise ImportError(
                "EPUB conversion requires 'ebooklib'. Install with: pip install ebooklib"
            ) from exc
        try:
            from bs4 import BeautifulSoup
        except ImportError as exc:
            raise ImportError(
                "EPUB conversion requires 'beautifulsoup4'. Install with: pip install beautifulsoup4"
            ) from exc
        try:
            from markdownify import markdownify as md
        except ImportError as exc:
            raise ImportError(
                "EPUB conversion requires 'markdownify'. Install with: pip install markdownify"
            ) from exc

        book = epub.read_epub(self.filepath)
        chapters = []

        def extract_item_content(item):
            soup = BeautifulSoup(item.get_body_content(), "html.parser")
            title = None
            if soup.title and soup.title.string:
                title = soup.title.string
            else:
                first_heading = soup.find(["h1", "h2", "h3"])
                if first_heading:
                    title = first_heading.get_text(strip=True)
            title = title or "Untitled"
            text_md = md(str(soup))
            return title.strip(), text_md.strip()

        def flatten_toc(toc):
            result = []
            if isinstance(toc, (list, tuple)):
                for x in toc:
                    result.extend(flatten_toc(x))
                return result

            if hasattr(toc, "href"):
                return [toc]

            return []

        toc_items = flatten_toc(book.toc)
        used_hrefs = set()
        for entry in toc_items:
            if not hasattr(entry, "href"):
                continue
            href = entry.href.split("#")[0]
            if href in used_hrefs:
                continue
            used_hrefs.add(href)
            item = book.get_item_with_href(href)
            if item:
                title, content = extract_item_content(item)
                chapters.append((title, content))

        if not chapters:
            for item in book.get_items_of_type(epub.ITEM_DOCUMENT):
                title, body = extract_item_content(item)
                chapters.append((title, body))

        return chapters

class DocxBook(Book):
    """A book in DOCX format."""

    def extract_chapters(self):
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "DOCX conversion requires 'python-docx'. Install with: pip install python-docx"
            ) from exc

        doc = Document(self.filepath)
        chapters = []
        current_title = "Untitled"
        current_content = []

        for para in doc.paragraphs:
            style = para.style.name.lower() if para.style and para.style.name else ""
            text = para.text.strip()
            if not text:
                continue

            if "heading" in style and ("1" in style or style == "heading"):
                if current_content:
                    chapters.append((current_title, "\n".join(current_content)))
                current_title = text
                current_content = []
            else:
                current_content.append(text)

        if current_content:
            chapters.append((current_title, "\n".join(current_content)))

        if not chapters:
            return [("Book", "\n".join([p.text for p in doc.paragraphs]))]
        return chapters

class PdfBook(Book):
    """A book in PDF format."""

    def extract_chapters(self):
        chapter_mode = self.options.get("chapter_mode", "auto")
        ocr_mode = self.options.get("ocr_mode", "auto")
        ocr_lang = self.options.get("ocr_lang", "chi_sim+eng")
        keep_intermediate = bool(self.options.get("keep_intermediate", False))
        chunk_words = int(self.options.get("chunk_words", 4000))
        toc_max_pages = int(self.options.get("toc_max_pages", 30))

        context = prepare_pdf_for_processing(
            self.filepath,
            ocr_mode=ocr_mode,
            ocr_lang=ocr_lang,
            keep_intermediate=keep_intermediate,
        )
        try:
            chapters = extract_pdf_chapters(
                context.processed_pdf,
                chapter_mode=chapter_mode,
                chunk_words=chunk_words,
                toc_max_pages=toc_max_pages,
            )
        except Exception as exc:
            note = finalize_pdf_processing(context, success=False, error=exc)
            if note:
                raise RuntimeError(f"{exc}\n{note}") from exc
            raise

        finalize_pdf_processing(context, success=True, error=None)
        return chapters

class TxtBook(Book):
    """A book in TXT format."""

    def extract_chapters(self):
        with open(self.filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        chapter_token = (
            r"(?:\d+|[IVXLCM]+|"
            r"ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|"
            r"ELEVEN|TWELVE|THIRTEEN|FOURTEEN|FIFTEEN|SIXTEEN|SEVENTEEN|EIGHTEEN|NINETEEN|TWENTY)"
        )
        pattern = re.compile(
            rf"(?P<title>^\s*(?:Chapter|Part|Book)\s+{chapter_token}\b.*$)",
            re.MULTILINE | re.IGNORECASE
        )
        matches = list(pattern.finditer(text))
        chapters = []

        if matches:
            for i in range(len(matches)):
                start = matches[i].end()
                end = matches[i + 1].start() if (i + 1) < len(matches) else len(text)
                title_line = matches[i].group("title").strip()
                content = text[start:end].strip()
                title = re.sub(r"^(Chapter|Part|Book)\s+", "", title_line, flags=re.IGNORECASE)
                chapters.append((title.strip() or f"Chapter {i+1}", content))
        else:
            chapters = [("Book", text)]
        return chapters

class BookFactory:
    """Factory for creating book objects."""

    @staticmethod
    def create_book(filepath, **options):
        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".epub":
            return EpubBook(filepath, **options)
        elif ext == ".docx":
            return DocxBook(filepath, **options)
        elif ext == ".pdf":
            return PdfBook(filepath, **options)
        elif ext == ".txt":
            return TxtBook(filepath, **options)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

class MarkdownWriter:
    """Writes chapters to Markdown files."""

    def __init__(self, outdir):
        self.outdir = outdir
        os.makedirs(self.outdir, exist_ok=True)

    def write_chapters(self, chapters):
        for i, (title, content) in enumerate(chapters, 1):
            safe_title = re.sub(r"[^a-zA-Z0-9_-]+", "_", title)[:50]
            filename = os.path.join(self.outdir, f"{i:02d}_{safe_title or 'Chapter'}.md")
            with open(filename, "w", encoding="utf-8") as f:
                f.write(f"# {title}\n\n{content.strip()}\n")

class BookConverter:
    """Orchestrates the book conversion process."""

    def __init__(
        self,
        input_file,
        outdir,
        chapter_mode="auto",
        ocr_mode="auto",
        ocr_lang="chi_sim+eng",
        keep_intermediate=False,
        chunk_words=4000,
        toc_max_pages=30,
    ):
        self.input_file = input_file
        self.outdir = outdir
        self.chapter_mode = chapter_mode
        self.ocr_mode = ocr_mode
        self.ocr_lang = ocr_lang
        self.keep_intermediate = keep_intermediate
        self.chunk_words = chunk_words
        self.toc_max_pages = toc_max_pages

    def convert(self):
        print(f"📘 Converting: {self.input_file}")
        book = BookFactory.create_book(
            self.input_file,
            chapter_mode=self.chapter_mode,
            ocr_mode=self.ocr_mode,
            ocr_lang=self.ocr_lang,
            keep_intermediate=self.keep_intermediate,
            chunk_words=self.chunk_words,
            toc_max_pages=self.toc_max_pages,
        )
        chapters = book.extract_chapters()
        print(f"→ Detected {len(chapters)} chapters. Writing to {self.outdir}/")
        writer = MarkdownWriter(self.outdir)
        writer.write_chapters(chapters)
        print("✅ Conversion complete.")
